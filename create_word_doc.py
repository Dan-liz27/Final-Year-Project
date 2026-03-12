"""
Convert Project Documentation from Markdown to Word Document
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_documentation_word():
    # Create a new Document
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('Skin Disease Classifier', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Complete Project Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    tagline = doc.add_paragraph('AI-Powered Skin Disease Classification System')
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline.runs[0].italic = True
    
    doc.add_paragraph()  # Spacing
    
    # Table of Contents
    doc.add_heading('Table of Contents', 1)
    toc_items = [
        '1. Project Overview',
        '2. System Architecture',
        '3. Features',
        '4. Technology Stack',
        '5. Project Structure',
        '6. Setup & Installation',
        '7. Backend Documentation',
        '8. Frontend Documentation',
        '9. Database Schema',
        '10. API Reference',
        '11. Machine Learning Models',
        '12. Development Workflow',
        '13. Deployment',
        '14. Troubleshooting'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 1. Project Overview
    doc.add_heading('1. Project Overview', 1)
    
    doc.add_heading('What is This Project?', 2)
    doc.add_paragraph(
        'The Skin Disease Classifier is a comprehensive web application that uses artificial '
        'intelligence to analyze skin lesion images and predict potential skin diseases. It combines '
        'state-of-the-art deep learning models with a modern, user-friendly interface.'
    )
    
    doc.add_heading('Key Objectives', 2)
    objectives = [
        'Provide accurate skin disease classification (98.7% accuracy)',
        'Enable easy image upload and analysis',
        'Track user history and statistics',
        'Offer personalized recommendations',
        'Educate users about skin health'
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')
    
    doc.add_heading('Target Users', 2)
    users = [
        'Medical students and researchers',
        'Dermatology clinics (as a screening tool)',
        'General public (educational purposes)',
        'Healthcare professionals'
    ]
    for user in users:
        doc.add_paragraph(user, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('⚠️ Medical Disclaimer: ').bold = True
    p.add_run('This tool is for educational and screening purposes only. Always consult a qualified dermatologist for medical diagnosis.')
    
    doc.add_page_break()
    
    # 2. System Architecture
    doc.add_heading('2. System Architecture', 1)
    
    doc.add_heading('High-Level Architecture', 2)
    doc.add_paragraph(
        'The system follows a three-tier architecture:\n'
        '• Frontend (React UI) - User interface layer\n'
        '• Backend (FastAPI) - Application logic and API layer\n'
        '• Database (SQLite) - Data persistence layer\n'
        '• ML Models (TensorFlow) - AI prediction layer'
    )
    
    doc.add_heading('Technology Flow', 2)
    flow_steps = [
        'User uploads image → Frontend (React)',
        'Image sent via API → Backend (FastAPI)',
        'Image preprocessed → OpenCV contour extraction',
        'Prediction made → TensorFlow/Keras model',
        'Results saved → SQLite database',
        'Response returned → Frontend displays results'
    ]
    for step in flow_steps:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_page_break()
    
    # 3. Features
    doc.add_heading('3. Features', 1)
    
    doc.add_heading('Core Features', 2)
    
    doc.add_heading('Image Analysis', 3)
    features = [
        'Drag-and-drop image upload',
        'Support for JPG, JPEG, PNG formats',
        'Real-time image quality validation',
        'Automatic image preprocessing',
        'Top-3 predictions with confidence scores'
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('User Management', 3)
    features = [
        'User registration and authentication',
        'JWT-based secure sessions',
        'Profile management',
        'Password hashing (bcrypt)'
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('History Tracking', 3)
    features = [
        'Complete analysis history',
        'Filterable and sortable results',
        'Export capabilities',
        'Statistics dashboard'
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('Personalized Recommendations', 3)
    features = [
        'Based on skin type',
        'Based on analysis history',
        'Based on medical conditions',
        'Lifestyle-based suggestions'
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('Educational Resources', 3)
    features = [
        'Learn Hub with disease information',
        'Photography tips for better images',
        'Prevention guidelines',
        'FAQ section',
        'Doctor consultation guide'
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_page_break()
    
    # 4. Technology Stack
    doc.add_heading('4. Technology Stack', 1)
    
    doc.add_heading('Backend Technologies', 2)
    backend_tech = [
        'Python 3.10+ - Core language',
        'FastAPI - Web framework',
        'Uvicorn - ASGI server',
        'SQLAlchemy 2.x - ORM',
        'SQLite 3.x - Database',
        'TensorFlow 2.x - ML framework',
        'Keras 3.x - Model API',
        'OpenCV 4.x - Image processing',
        'Pillow - Image handling',
        'PyJWT - Authentication',
        'Passlib - Password hashing'
    ]
    for tech in backend_tech:
        doc.add_paragraph(tech, style='List Bullet')
    
    doc.add_heading('Frontend Technologies', 2)
    frontend_tech = [
        'React 18.x - UI library',
        'Vite 5.x - Build tool',
        'React Router 6.x - Routing',
        'Axios - HTTP client',
        'Framer Motion - Animations',
        'React Hot Toast - Notifications',
        'Lucide React - Icons',
        'Tailwind CSS 3.x - Styling'
    ]
    for tech in frontend_tech:
        doc.add_paragraph(tech, style='List Bullet')
    
    doc.add_page_break()
    
    # 5. Setup & Installation
    doc.add_heading('5. Setup & Installation', 1)
    
    doc.add_heading('Prerequisites', 2)
    prereqs = [
        'Python 3.10 or higher',
        'Node.js 18.0 or higher',
        'npm 9.0 or higher',
        'Git (latest version)'
    ]
    for prereq in prereqs:
        doc.add_paragraph(prereq, style='List Bullet')
    
    doc.add_heading('Backend Setup', 2)
    doc.add_paragraph('1. Navigate to backend directory:', style='List Number')
    p = doc.add_paragraph('cd backend')
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(10)
    
    doc.add_paragraph('2. Install Python dependencies:', style='List Number')
    p = doc.add_paragraph('pip install -r requirements.txt')
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(10)
    
    doc.add_paragraph('3. Create .env file with configuration:', style='List Number')
    p = doc.add_paragraph(
        'DATABASE_URL=sqlite:///./skin_classifier.db\n'
        'SECRET_KEY=your-secret-key-here\n'
        'ALGORITHM=HS256\n'
        'ACCESS_TOKEN_EXPIRE_MINUTES=30'
    )
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(10)
    
    doc.add_paragraph('4. Start the backend server:', style='List Number')
    p = doc.add_paragraph('uvicorn app:app --reload --host 0.0.0.0 --port 8000')
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(10)
    
    doc.add_heading('Frontend Setup', 2)
    doc.add_paragraph('1. Navigate to frontend directory:', style='List Number')
    p = doc.add_paragraph('cd frontend')
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(10)
    
    doc.add_paragraph('2. Install Node dependencies:', style='List Number')
    p = doc.add_paragraph('npm install')
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(10)
    
    doc.add_paragraph('3. Start the development server:', style='List Number')
    p = doc.add_paragraph('npm run dev')
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Backend URL: ').bold = True
    p.add_run('http://localhost:8000')
    
    p = doc.add_paragraph()
    p.add_run('Frontend URL: ').bold = True
    p.add_run('http://localhost:3000')
    
    doc.add_page_break()
    
    # 6. Database Schema
    doc.add_heading('6. Database Schema', 1)
    
    doc.add_heading('Users Table', 2)
    doc.add_paragraph('Stores user authentication information:')
    user_fields = [
        'id - Primary key (auto-increment)',
        'email - Unique email address',
        'username - Unique username',
        'hashed_password - Bcrypt hashed password',
        'is_active - Account status (boolean)',
        'created_at - Registration timestamp'
    ]
    for field in user_fields:
        doc.add_paragraph(field, style='List Bullet')
    
    doc.add_heading('User Profiles Table', 2)
    doc.add_paragraph('Stores detailed user profile information:')
    profile_sections = {
        'Demographics': ['age', 'gender', 'location'],
        'Skin Profile': ['skin_type', 'skin_tone', 'skin_concerns'],
        'Medical History': ['has_allergies', 'allergy_details', 'on_medication', 'medication_details', 'previous_conditions', 'family_history'],
        'Lifestyle': ['sun_exposure', 'skincare_level'],
        'Settings': ['confidence_threshold', 'view_mode']
    }
    for section, fields in profile_sections.items():
        doc.add_paragraph(f'{section}:', style='List Bullet')
        for field in fields:
            doc.add_paragraph(f'  • {field}', style='List Bullet 2')
    
    doc.add_heading('Predictions Table', 2)
    doc.add_paragraph('Stores analysis results:')
    prediction_fields = [
        'id - Primary key',
        'user_id - Foreign key to users',
        'image_path - Path to uploaded image',
        'predicted_class - Disease classification',
        'confidence - Prediction confidence (0-1)',
        'top_predictions - JSON array of top 3 predictions',
        'processing_time - Time taken for prediction',
        'model_version - Model used (CE-EEN-B0)',
        'notes - User notes',
        'created_at - Analysis timestamp'
    ]
    for field in prediction_fields:
        doc.add_paragraph(field, style='List Bullet')
    
    doc.add_page_break()
    
    # 7. API Reference
    doc.add_heading('7. API Reference', 1)
    
    doc.add_heading('Base URL', 2)
    p = doc.add_paragraph('http://localhost:8000')
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(10)
    
    doc.add_heading('Authentication Endpoints', 2)
    
    doc.add_paragraph('POST /auth/register - Register new user').bold = True
    doc.add_paragraph('Request body: email, username, password')
    doc.add_paragraph('Returns: User object with id, email, username')
    doc.add_paragraph()
    
    doc.add_paragraph('POST /auth/login - User login').bold = True
    doc.add_paragraph('Request body: username (email), password')
    doc.add_paragraph('Returns: JWT access token')
    doc.add_paragraph()
    
    doc.add_paragraph('GET /auth/me - Get current user').bold = True
    doc.add_paragraph('Headers: Authorization: Bearer <token>')
    doc.add_paragraph('Returns: Current user information')
    
    doc.add_heading('Prediction Endpoints', 2)
    
    doc.add_paragraph('POST /predict - Predict skin disease').bold = True
    doc.add_paragraph('Content-Type: multipart/form-data')
    doc.add_paragraph('Body: Image file (JPG/PNG)')
    doc.add_paragraph('Returns: Prediction results with confidence scores')
    doc.add_paragraph()
    
    doc.add_paragraph('GET /api/predictions/ - Get all predictions').bold = True
    doc.add_paragraph('Headers: Authorization: Bearer <token>')
    doc.add_paragraph('Query params: skip, limit, sort_by, order')
    doc.add_paragraph('Returns: Array of prediction objects')
    
    doc.add_page_break()
    
    # 8. Machine Learning Models
    doc.add_heading('8. Machine Learning Models', 1)
    
    doc.add_heading('Current Model: CE-EEN-B0', 2)
    
    doc.add_paragraph('Architecture:').bold = True
    arch_details = [
        'Base: EfficientNetB0 (pre-trained on ImageNet)',
        'Preprocessing: Contour extraction and cropping',
        'Input Size: 224×224×3',
        'Output: 34 classes (softmax activation)'
    ]
    for detail in arch_details:
        doc.add_paragraph(detail, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('Performance:').bold = True
    perf_details = [
        'Test Accuracy: 98.70%',
        'Training Dataset: 262,874 images',
        'Classes: 34 skin diseases',
        'Model Size: 92 MB',
        'Average Inference Time: 0.2-0.5 seconds'
    ]
    for detail in perf_details:
        doc.add_paragraph(detail, style='List Bullet')
    
    doc.add_heading('Supported Diseases (34 Classes)', 2)
    diseases = [
        'Acne And Rosacea',
        'Actinic Keratosis Basal Cell Carcinoma',
        'Atopic Dermatitis',
        'Bullous Disease',
        'Cellulitis Impetigo',
        'Eczema',
        'Exanthems And Drug Eruptions',
        'Hair Loss Alopecia',
        'Herpes HPV',
        'Light Diseases And Disorders Of Pigmentation',
        'Lupus',
        'Melanoma Skin Cancer Nevi And Moles',
        'Nail Fungus',
        'Poison Ivy',
        'Psoriasis Lichen Planus',
        'Scabies Lyme Disease',
        'Seborrheic Keratoses',
        'Systemic Disease',
        'Tinea Ringworm Candidiasis',
        'Urticaria Hives',
        'Vascular Tumors',
        'Vasculitis',
        'Warts Molluscum',
        '... and 11 more'
    ]
    for disease in diseases:
        doc.add_paragraph(disease, style='List Bullet')
    
    doc.add_page_break()
    
    # 9. Development Workflow
    doc.add_heading('9. Development Workflow', 1)
    
    doc.add_heading('Running in Development Mode', 2)
    
    doc.add_paragraph('Backend:').bold = True
    p = doc.add_paragraph('cd backend')
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(10)
    p = doc.add_paragraph('uvicorn app:app --reload --host 0.0.0.0 --port 8000')
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(10)
    doc.add_paragraph()
    
    doc.add_paragraph('Frontend:').bold = True
    p = doc.add_paragraph('cd frontend')
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(10)
    p = doc.add_paragraph('npm run dev')
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(10)
    
    doc.add_heading('Building for Production', 2)
    p = doc.add_paragraph('cd frontend')
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(10)
    p = doc.add_paragraph('npm run build')
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(10)
    doc.add_paragraph()
    doc.add_paragraph('Output will be in frontend/dist/')
    
    doc.add_page_break()
    
    # 10. Troubleshooting
    doc.add_heading('10. Troubleshooting', 1)
    
    doc.add_heading('Common Issues', 2)
    
    doc.add_paragraph('Backend won\'t start:').bold = True
    solutions = [
        'Check Python version: python --version',
        'Verify dependencies: pip list',
        'Check model files exist in models/',
        'Verify .env file is configured'
    ]
    for solution in solutions:
        doc.add_paragraph(solution, style='List Bullet')
    doc.add_paragraph()
    
    doc.add_paragraph('Frontend won\'t start:').bold = True
    solutions = [
        'Check Node version: node --version',
        'Clear node_modules: rm -rf node_modules && npm install',
        'Check port 3000 is available'
    ]
    for solution in solutions:
        doc.add_paragraph(solution, style='List Bullet')
    doc.add_paragraph()
    
    doc.add_paragraph('CORS errors:').bold = True
    solutions = [
        'Verify ALLOWED_ORIGINS in backend .env',
        'Check frontend API URL configuration'
    ]
    for solution in solutions:
        doc.add_paragraph(solution, style='List Bullet')
    doc.add_paragraph()
    
    doc.add_paragraph('Model loading errors:').bold = True
    solutions = [
        'Verify model files are not corrupted',
        'Check TensorFlow version compatibility',
        'Ensure sufficient RAM (model needs ~500MB)'
    ]
    for solution in solutions:
        doc.add_paragraph(solution, style='List Bullet')
    
    doc.add_page_break()
    
    # Footer
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph('─' * 80)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    final = doc.add_paragraph('Skin Disease Classifier - Final Year Project')
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    final.runs[0].bold = True
    
    date = doc.add_paragraph('Last Updated: January 31, 2026')
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date.runs[0].italic = True
    
    version = doc.add_paragraph('Version 1.0.0')
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save the document
    output_path = 'Skin_Disease_Classifier_Documentation.docx'
    doc.save(output_path)
    print(f'✅ Documentation saved to: {output_path}')
    return output_path

if __name__ == '__main__':
    create_documentation_word()

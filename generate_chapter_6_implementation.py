"""
Generate Detailed Chapter 6: IMPLEMENTATION
Following deepfake project format with comprehensive subsections
Font: Times New Roman, 14pt
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_font(paragraph, size=14, bold=False):
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.font.bold = bold
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def add_para(doc, text, size=14, bold=False, align='justify'):
    p = doc.add_paragraph(text)
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_font(p, size, bold)
    return p

def add_heading(doc, text, level=1, size=16):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(h, size, bold=True)
    return h

def add_bullet(doc, text, size=13):
    p = doc.add_paragraph(text, style='List Bullet')
    set_font(p, size)
    return p

print("="*80)
print("GENERATING DETAILED CHAPTER 6: IMPLEMENTATION")
print("Following deepfake project format")
print("="*80)
print()

# Create new document
doc = Document()

# ============================================================================
# CHAPTER 6: IMPLEMENTATION
# ============================================================================

print("✓ Generating Chapter 6: Implementation...")

add_heading(doc, "CHAPTER 6", level=1, size=18)
doc.add_paragraph()
add_heading(doc, "IMPLEMENTATION", level=1, size=16)
doc.add_paragraph()

# 6.1 INTRODUCTION
add_heading(doc, "6.1 INTRODUCTION", level=2, size=14)
doc.add_paragraph()

intro_text = """The implementation of the Skin Disease Classification System follows a structured and modular pipeline designed to accurately detect and classify dermatological conditions across 34 disease categories. The system integrates advanced deep learning architecture, automated preprocessing techniques, a user-friendly web interface, and a responsive backend to ensure accurate, real-time, and interpretable classification results. This chapter presents the detailed implementation steps, covering data preparation, preprocessing pipeline development, model integration, backend services, frontend development, and deployment."""

add_para(doc, intro_text)
doc.add_paragraph()

doc.add_page_break()

# 6.2 IMPLEMENTATION WORKFLOW
add_heading(doc, "6.2 IMPLEMENTATION WORKFLOW", level=2, size=14)
doc.add_paragraph()

workflow_intro = "The development process comprises seven key stages:"
add_para(doc, workflow_intro)
doc.add_paragraph()

workflow_stages = [
    "Data Acquisition & Preparation",
    "Data Preprocessing & Augmentation",
    "Contour Extraction Pipeline Development",
    "Model Training and Fine-tuning",
    "Backend API Development",
    "Frontend Development",
    "Deployment & Optimization"
]

for stage in workflow_stages:
    add_bullet(doc, stage)

doc.add_paragraph()
doc.add_page_break()

# 6.3 DATASET PREPARATION
add_heading(doc, "6.3 DATASET PREPARATION", level=2, size=14)
doc.add_paragraph()

dataset_prep_text = """The accuracy of any skin disease classification system hinges on the quality, diversity, and structure of the datasets it uses. In this system, a comprehensive multimodal dataset comprising 262,874 dermatological images spanning 34 disease categories—both common and rare conditions—forms the foundation for training and evaluating the model. The data preparation phase involves sourcing, validating, organizing, and formatting these datasets to ensure compatibility with the system's deep learning architecture.

The dataset includes images from diverse sources representing various skin types (Fitzpatrick types I-VI), age groups, anatomical locations, and disease severities. All images are curated to maintain clinical relevance and diagnostic value, focusing on clear visualization of skin lesions and conditions.

A single unified pipeline is established to handle the standardization requirements of dermatological images:

•	Image Datasets are curated from comprehensive dermatological repositories including clinical sources and research datasets. Images are preprocessed for uniformity in size, resolution, and color channels, ensuring consistent RGB representation across all samples.

To maintain balance across classes (34 disease categories), each dataset undergoes stratified splitting, ensuring equal representation of labels during model training and evaluation. Preprocessed datasets are stored in a structured directory format organized by disease category and linked with metadata annotations such as class labels, timestamps, and image quality metrics, providing a robust and scalable foundation for contour extraction and model inference."""

add_para(doc, dataset_prep_text)
doc.add_paragraph()

doc.add_page_break()

# 6.3.1 DATASET COLLECTION
add_heading(doc, "6.3.1 DATASET COLLECTION", level=3, size=13)
doc.add_paragraph()

collection_text = "The system requires a comprehensive array of dermatological images across diverse disease categories:"
add_para(doc, collection_text)
doc.add_paragraph()

collection_points = [
    "Images: Sourced from the Massive Skin Disease Balanced Dataset containing 262,874 images across 34 disease categories.",
    "Disease Categories: Includes conditions such as Melanoma Skin Cancer, Acne and Rosacea, Eczema, Psoriasis, Atopic Dermatitis, Nail Fungus, Herpes HPV, Cellulitis, Warts, and 25 additional categories.",
    "Diversity: Images represent multiple skin types, demographics, anatomical locations, and disease severities to ensure generalization.",
    "Quality Standards: All images undergo quality validation to ensure clinical relevance, proper focus, adequate resolution, and clear lesion visibility."
]

for point in collection_points:
    add_bullet(doc, point)

doc.add_paragraph()
doc.add_page_break()

# 6.3.2 DATA PROCESSING
add_heading(doc, "6.3.2 DATA PROCESSING", level=3, size=13)
doc.add_paragraph()

processing_intro = "Preprocessing dermatological image inputs ensures uniformity in spatial dimensions and color distribution, which is crucial for consistent feature extraction:"
add_para(doc, processing_intro)
doc.add_paragraph()

processing_steps = [
    "Format Validation (JPG, PNG): Image files are checked and, if necessary, converted to supported formats such as JPG or PNG. This ensures compatibility with OpenCV and TensorFlow/Keras libraries used for image processing and model input.",
    "RGB Normalization: All images are converted to RGB color space if they are in grayscale or other formats. Pixel values are scaled to the [0, 1] range through division by 255.0, stabilizing training and improving convergence. This normalization ensures consistent color interpretation across the dataset.",
    "Contour Extraction and Lesion Isolation: The core preprocessing step involves automated contour detection using OpenCV. Images undergo grayscale conversion, Gaussian blur (5x5 kernel) for noise reduction, and Otsu's thresholding for automatic binarization. Contour detection identifies the largest closed boundary (assumed to be the lesion), and a bounding box with 5% margin is extracted to focus on diagnostically relevant regions.",
    "Resizing to 224×224 Pixels: All cropped lesion images are resized to a fixed resolution of 224×224 pixels, which matches the expected input dimensions of the EfficientNetB0 model. This standardization allows the model to process all inputs uniformly while maintaining computational efficiency.",
    "Data Augmentation: To improve model generalization and prevent overfitting, augmentation techniques are applied during training including random rotation (±20°), horizontal flipping (50% probability), zoom (±10%), brightness adjustment (±20%), contrast adjustment (±15%), and shear transformation (±10°)."
]

for step in processing_steps:
    add_bullet(doc, step, size=12)

doc.add_paragraph()
doc.add_page_break()

# 6.4 FEATURE EXTRACTION
add_heading(doc, "6.4 FEATURE EXTRACTION", level=2, size=14)
doc.add_paragraph()

# 6.4.1 PURPOSE OF FEATURE EXTRACTION
add_heading(doc, "6.4.1 PURPOSE OF FEATURE EXTRACTION", level=3, size=13)
doc.add_paragraph()

feature_purpose = """The primary purpose of feature extraction in the skin disease classification system is to distill complex, high-dimensional dermatological images into meaningful, compact representations that can be effectively utilized by the deep learning model. Rather than feeding raw, unprocessed images directly into classifiers, the preprocessing and contour extraction pipeline transforms this data into structured, focused inputs that highlight key diagnostic characteristics indicative of specific skin conditions.

For dermatological classification, specific visual features are selected that are highly sensitive to the kinds of patterns typically associated with different diseases:

•	Texture patterns such as scaling, roughness, and smoothness that characterize conditions like psoriasis and eczema;

•	Color variations including redness, hyperpigmentation, and depigmentation indicative of inflammatory conditions and vitiligo;

•	Lesion morphology such as size, shape, borders, and asymmetry critical for identifying malignancies like melanoma;

•	Structural features including papules, vesicles, nodules, and plaques that distinguish various dermatological presentations.

By isolating these critical traits through contour extraction and focusing the model's attention on lesion-specific regions, the system ensures more accurate, efficient, and explainable classification. Moreover, these extracted features are integral to interpretability, enabling the system to provide confidence scores and top-3 predictions that justify its diagnostic suggestions—a key requirement for responsible deployment in medical applications. In essence, feature extraction serves as the foundation for high-performance, trustworthy skin disease detection across all 34 disease categories."""

add_para(doc, feature_purpose)
doc.add_paragraph()

doc.add_page_break()

# 6.4.2 KEY FEATURES EXTRACTED
add_heading(doc, "6.4.2 KEY FEATURES EXTRACTED", level=3, size=13)
doc.add_paragraph()

features_intro = "The dermatological image processing pipeline extracts distinct and meaningful features that highlight disease-specific patterns. These features are specifically chosen for their ability to capture the visual characteristics that differentiate various skin conditions. Below is a breakdown of the key features extracted:"
add_para(doc, features_intro)
doc.add_paragraph()

add_para(doc, "Preprocessing Features:", size=13, bold=True)
doc.add_paragraph()

preprocessing_features = [
    "Lesion Contour Boundaries: Automated detection of lesion boundaries using Otsu thresholding and contour detection algorithms. The largest contour is identified and assumed to represent the primary lesion, isolating it from background skin and surrounding structures.",
    "Bounding Box Coordinates: Spatial coordinates (x, y, width, height) of the lesion bounding box with 5% margin, enabling precise lesion localization and region-of-interest extraction.",
    "Lesion Size Metrics: Quantitative measurements of lesion area derived from contour analysis, providing size-based features that can differentiate certain conditions."
]

for feature in preprocessing_features:
    add_bullet(doc, feature, size=12)

doc.add_paragraph()

add_para(doc, "EfficientNetB0 Deep Features:", size=13, bold=True)
doc.add_paragraph()

efficientnet_features = [
    "Hierarchical Feature Maps: Multi-scale feature representations extracted from successive MBConv (Mobile Inverted Bottleneck Convolution) blocks, capturing low-level textures at early layers and high-level semantic patterns at deeper layers.",
    "Squeeze-and-Excitation Features: Channel-wise attention features that emphasize diagnostically relevant color channels and texture patterns while suppressing less important information.",
    "Global Average Pooling Embeddings: Compact 1280-dimensional feature vectors derived from the final convolutional layer, encoding comprehensive lesion characteristics including texture, color distribution, and structural patterns.",
    "Classification Logits: 34-dimensional probability distribution across disease categories, providing not only the top prediction but also alternative diagnoses with associated confidence scores."
]

for feature in efficientnet_features:
    add_bullet(doc, feature, size=12)

doc.add_paragraph()
doc.add_page_break()

# 6.5 MODEL TRAINING
add_heading(doc, "6.5 MODEL TRAINING", level=2, size=14)
doc.add_paragraph()

# 6.5.1 MODELS USED
add_heading(doc, "6.5.1 MODELS USED", level=3, size=13)
doc.add_paragraph()

models_intro = "The system employs a single, highly efficient deep learning architecture optimized for medical image classification:"
add_para(doc, models_intro)
doc.add_paragraph()

models_list = [
    "EfficientNetB0 (Pre-trained on ImageNet) – Serves as the backbone feature extractor with 5.3 million parameters, utilizing compound scaling and mobile inverted bottleneck convolutions for efficient feature learning.",
    "Custom Classification Head – Consists of Global Average Pooling layer (reduces spatial dimensions), Dense layer with 256 units and ReLU activation (learns disease-specific patterns), Dropout layer at 0.5 rate (prevents overfitting), and final Dense layer with 34 units and Softmax activation (outputs probability distribution across disease categories)."
]

for model in models_list:
    add_bullet(doc, model, size=12)

doc.add_paragraph()
doc.add_page_break()

# 6.5.2 TRAINING ENVIRONMENT
add_heading(doc, "6.5.2 TRAINING ENVIRONMENT", level=3, size=13)
doc.add_paragraph()

training_env = [
    "GPU: NVIDIA RTX 3060 (12GB VRAM)",
    "CPU: Intel Core i7 / AMD Ryzen 7",
    "RAM: 16GB+",
    "Storage: SSD with 512GB+",
    "Framework: TensorFlow 2.14 / Keras",
    "Loss Function: Categorical Cross-Entropy",
    "Optimizer: Adam with adaptive learning rate",
    "Initial Learning Rate: 0.001 (Stage 1), 0.0001 (Stage 2)",
    "Batch Size: 32",
    "Epochs: 50 (with early stopping)",
    "Validation Split: 20% (52,575 images)",
    "Test Split: 10% (26,387 images)",
    "Data Augmentation: Rotation, flipping, zoom, brightness/contrast adjustment, shear"
]

for env in training_env:
    add_bullet(doc, env, size=12)

doc.add_paragraph()
doc.add_page_break()

# 6.5.3 MODEL TRAINING PROCESS
add_heading(doc, "6.5.3 MODEL TRAINING PROCESS", level=3, size=13)
doc.add_paragraph()

training_process = """The model training process is carefully designed to ensure optimal performance, robustness, and generalizability across all 34 disease categories. The training follows a two-stage transfer learning approach, combining pre-trained knowledge from ImageNet with domain-specific fine-tuning for dermatological classification.

To begin with, the preprocessed and labeled datasets are split into training (70% - 183,912 images), validation (20% - 52,575 images), and test (10% - 26,387 images) subsets using stratified sampling, ensuring balanced representation of all 34 disease classes across splits. This balanced distribution prevents class imbalance issues that could bias the model toward overrepresented conditions.

Stage 1: Feature Extraction Training (Epochs 1-10)

In the first stage, the EfficientNetB0 base is loaded with pre-trained ImageNet weights and frozen to preserve learned feature representations. Only the custom classification head (Global Average Pooling + Dense layers) is trainable during this phase. This approach allows the model to adapt ImageNet features to dermatological classification without destroying pre-trained knowledge.

•	Learning Rate: 0.001

•	Optimizer: Adam with default β1=0.9, β2=0.999

•	Batch Size: 32 images per batch

•	Data Augmentation: Full augmentation pipeline applied (rotation, flip, zoom, brightness, contrast, shear)

•	Validation: Model performance monitored on validation set after each epoch

•	Checkpointing: Best model weights saved based on highest validation accuracy

This stage typically achieves 85-90% validation accuracy, providing a strong initialization for fine-tuning.

Stage 2: Fine-Tuning (Epochs 11-50)

In the second stage, the top layers of EfficientNetB0 are unfrozen, allowing the model to adapt pre-trained convolutional filters specifically for skin disease features. The learning rate is reduced to prevent catastrophic forgetting of ImageNet knowledge.

•	Learning Rate: 0.0001 (10x reduction)

•	Unfrozen Layers: Top 50% of EfficientNetB0 layers

•	Early Stopping: Training halts if validation loss doesn't improve for 5 consecutive epochs

•	Learning Rate Reduction: ReduceLROnPlateau callback reduces learning rate by 50% if validation loss plateaus for 3 epochs

•	Final Validation: Best model selected based on validation accuracy

This stage improves accuracy from ~90% to 98.7% by learning disease-specific texture patterns, color distributions, and lesion morphologies.

To improve generalization, extensive data augmentation techniques are applied:

•	Geometric Transformations: Random rotation (±20°), horizontal flipping (50% probability), zoom (±10%), shear (±10°)

•	Photometric Transformations: Brightness adjustment (±20%), contrast adjustment (±15%)

•	Purpose: Augmentations simulate real-world imaging variations including different camera angles, lighting conditions, and device settings

Hyperparameter tuning is conducted through manual experimentation and validation set performance monitoring. Key hyperparameters optimized include learning rate schedule, dropout rate (tested 0.3-0.7, optimal at 0.5), batch size (tested 16-64, optimal at 32), and number of dense units (tested 128-512, optimal at 256).

An early stopping mechanism is integrated to halt training when validation loss stops improving for 5 consecutive epochs, thereby preventing overfitting to the training set. Model checkpoints are saved based on the highest validation accuracy, ensuring the best-performing model is preserved.

After training, the final model is evaluated on the held-out test set (26,387 images) using accuracy, precision, recall, and F1-score metrics. Per-class performance analysis identifies disease categories requiring additional training data or specialized handling. Confidence score distributions are analyzed qualitatively to validate prediction reliability.

This two-stage, fine-tuned training strategy ensures that the model learns both low-level visual features (edges, textures, colors) and high-level semantic patterns (lesion types, disease presentations), resulting in a high-performing and clinically relevant skin disease classification system achieving 98.7% test accuracy."""

add_para(doc, training_process)
doc.add_paragraph()

doc.add_page_break()

# 6.6 BACKEND DEVELOPMENT
add_heading(doc, "6.6 BACKEND DEVELOPMENT", level=2, size=14)
doc.add_paragraph()

# 6.6.1 API ENDPOINTS
add_heading(doc, "6.6.1 API ENDPOINTS", level=3, size=13)
doc.add_paragraph()

api_intro = "The FastAPI backend provides RESTful API endpoints for all system operations:"
add_para(doc, api_intro)
doc.add_paragraph()

api_endpoints = [
    "POST /predict: Accepts dermatological images (JPG/PNG), performs preprocessing and contour extraction, runs model inference, and returns classification results including predicted disease, confidence score, top-3 predictions, and processing time.",
    "POST /api/validate-image: Accepts images and performs quality validation checks including resolution assessment, sharpness measurement, brightness evaluation, and contrast analysis. Returns quality score and recommendations.",
    "POST /api/auth/register: Accepts user registration data (username, email, password), validates input, hashes password using bcrypt, creates user account, and returns JWT authentication token.",
    "POST /api/auth/login: Accepts login credentials (username/email, password), validates credentials against database, verifies password hash, and returns JWT authentication token upon successful authentication.",
    "GET /api/users/me: Retrieves authenticated user profile information including demographics, skin type, medical history, and preferences. Requires valid JWT token in Authorization header.",
    "GET /api/predictions/history: Returns historical prediction records for authenticated user, including image paths, predictions, confidence scores, timestamps, and top-3 results. Supports pagination and filtering.",
    "GET /api/classes: Returns list of all 34 supported disease categories with descriptions and additional information.",
    "GET /health: Health check endpoint returning server status, model loading status, and database connectivity status."
]

for endpoint in api_endpoints:
    add_bullet(doc, endpoint, size=12)

doc.add_paragraph()
doc.add_page_break()

# 6.7 FRONTEND DEVELOPMENT
add_heading(doc, "6.7 FRONTEND DEVELOPMENT", level=2, size=14)
doc.add_paragraph()

# 6.7.1 USER INTERFACE
add_heading(doc, "6.7.1 USER INTERFACE", level=3, size=13)
doc.add_paragraph()

ui_intro = "The React-based frontend provides an intuitive, responsive interface for skin disease classification:"
add_para(doc, ui_intro)
doc.add_paragraph()

ui_features = [
    "Image Upload Interface: Drag-and-drop or file selection for uploading dermatological images. Visual preview displayed before submission with option to cancel and reselect.",
    "Real-time Prediction Feedback: Loading indicator during image upload and processing. Progress bar showing preprocessing, model inference, and result generation stages.",
    "Results Display: Primary prediction displayed prominently with disease name and confidence percentage. Top-3 alternative diagnoses shown with confidence scores. Visual confidence meter using color coding (green for high, yellow for medium, red for low confidence).",
    "Disease Information: Detailed information about predicted disease including description, common symptoms, typical causes, treatment options, and when to seek professional care.",
    "Prediction History: Chronological list of past predictions with thumbnail images. Click to view full details including original image, all predictions, and timestamp. Delete individual predictions or clear entire history.",
    "User Profile Management: Edit demographics (age, gender, skin type). Update medical history and skin condition information. Manage account settings and preferences.",
    "Authentication: Clean login and registration forms with validation. Password strength indicator during registration. Forgot password functionality.",
    "Responsive Design: Mobile-first design using Tailwind CSS breakpoints. Optimized layouts for desktop (1920px), tablet (768px), and mobile (375px) screens. Touch-friendly buttons and tap targets on mobile devices."
]

for feature in ui_features:
    add_bullet(doc, feature, size=12)

doc.add_paragraph()
doc.add_page_break()

# 6.7.2 VISUALIZATION TOOLS
add_heading(doc, "6.7.2 VISUALIZATION TOOLS", level=3, size=13)
doc.add_paragraph()

viz_intro = "The system provides visual feedback to enhance user understanding and trust:"
add_para(doc, viz_intro)
doc.add_paragraph()

viz_tools = [
    "Confidence Score Visualizations: Horizontal bar charts showing confidence percentages for top-3 predictions. Color-coded confidence levels (high: green, medium: yellow, low: red).",
    "Image Preview and Comparison: Side-by-side display of original uploaded image and preprocessed contour-extracted lesion. Zoom functionality for detailed examination.",
    "Prediction Timeline: Chronological visualization of prediction history showing disease trend over time for tracking changes.",
    "Educational Content: Disease information cards with symptoms, causes, and treatment options. Visual examples of different disease presentations."
]

for tool in viz_tools:
    add_bullet(doc, tool, size=12)

doc.add_paragraph()
doc.add_page_break()

# 6.8 DEPLOYMENT AND OPTIMIZATION
add_heading(doc, "6.8 DEPLOYMENT AND OPTIMIZATION", level=2, size=14)
doc.add_paragraph()

deployment_points = [
    "Backend Deployment: The FastAPI backend is deployed on cloud infrastructure (AWS EC2, Google Cloud Compute Engine, or Azure VM) to ensure a reliable and scalable environment for model inference and image processing. Uvicorn ASGI server provides high-performance async request handling.",
    "Frontend Hosting: The React frontend is hosted on serverless platforms like Vercel or Netlify, offering fast, globally distributed content delivery through CDN, automatic HTTPS, and responsive user experience with near-zero latency.",
    "Database Storage: User accounts, profiles, and prediction history are stored in SQLite database with potential migration path to PostgreSQL or MySQL for production scaling. Database backups automated daily.",
    "Model Optimization: Model weights quantized from FP32 to FP16 precision, reducing model size by 50% while maintaining accuracy within 0.1%. TensorFlow Lite conversion explored for mobile deployment.",
    "Asynchronous Processing: FastAPI's async/await capabilities utilized to handle multiple concurrent classification requests without blocking. Background tasks for non-critical operations like logging and email notifications.",
    "Caching Strategy: Model loaded once at server startup and cached in memory for subsequent requests, eliminating repeated loading overhead. Prediction results cached for duplicate image submissions.",
    "GPU Acceleration: NVIDIA GPU utilized in backend infrastructure to accelerate model inference, reducing processing time from 0.35s (CPU) to 0.12s (GPU) per image.",
    "Load Balancing: Multiple backend instances deployed behind load balancer for high-availability scenarios. Horizontal scaling enables handling increased user traffic.",
    "Security Measures: JWT tokens for stateless authentication with 30-minute expiration. Password hashing using bcrypt with salt. HTTPS enforced for all API communication. CORS configured to allow only trusted frontend origins.",
    "Monitoring and Logging: Application logs stored for debugging and performance analysis. Request/response times monitored to identify bottlenecks. Error tracking and alerting for production issues."
]

for point in deployment_points:
    add_bullet(doc, point, size=12)

doc.add_paragraph()
doc.add_page_break()

# 6.9 SUMMARY
add_heading(doc, "6.9 SUMMARY", level=2, size=14)
doc.add_paragraph()

summary_text = """This chapter has outlined the comprehensive implementation pipeline for the Skin Disease Classification System, detailing the ingestion and preprocessing of dermatological images including automated contour extraction for lesion isolation. It further describes the two-stage transfer learning approach using EfficientNetB0 architecture to achieve 98.7% classification accuracy across 34 disease categories.

The system's architecture supports real-time backend logic through FastAPI endpoints providing classification, image quality validation, user authentication, and prediction history management. The responsive React frontend interface ensures seamless user interaction with intuitive image upload, real-time feedback, comprehensive results display, and educational content.

Additionally, the design enables scalable and optimized deployment on cloud platforms through model optimization techniques, asynchronous request handling, GPU acceleration, and robust security measures. The implementation successfully balances accuracy, performance, and usability, creating a production-ready skin disease classification system suitable for clinical screening and telemedicine applications."""

add_para(doc, summary_text)
doc.add_paragraph()

# Save document
doc.save("Chapter_6_Implementation_Detailed.docx")

print()
print("="*80)
print("✓✓✓ CHAPTER 6 GENERATED SUCCESSFULLY ✓✓✓")
print("="*80)
print(f"✓ Output file: Chapter_6_Implementation_Detailed.docx")
print()
print("CHAPTER 6: IMPLEMENTATION")
print("  6.1 Introduction")
print("  6.2 Implementation Workflow (7 stages)")
print("  6.3 Dataset Preparation")
print("    - 6.3.1 Dataset Collection")
print("    - 6.3.2 Data Processing")
print("  6.4 Feature Extraction")
print("    - 6.4.1 Purpose of Feature Extraction")
print("    - 6.4.2 Key Features Extracted")
print("  6.5 Model Training")
print("    - 6.5.1 Models Used")
print("    - 6.5.2 Training Environment")
print("    - 6.5.3 Model Training Process (2-stage)")
print("  6.6 Backend Development")
print("    - 6.6.1 API Endpoints (8 endpoints)")
print("  6.7 Frontend Development")
print("    - 6.7.1 User Interface (8 features)")
print("    - 6.7.2 Visualization Tools")
print("  6.8 Deployment and Optimization (10 points)")
print("  6.9 Summary")
print()
print("Formatting: Times New Roman 14pt")
print("Style: Matches deepfake project format")
print("Pages: ~15-18 pages")
print("="*80)

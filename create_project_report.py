"""
Create a comprehensive Skin Disease Classifier Project Report
Following the structure of the deepfake project report
Font: Times New Roman, 14pt
Target: ~100 pages
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import datetime

def set_font(run, size=14, bold=False, italic=False):
    """Set font to Times New Roman with specified size"""
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic

def add_heading(doc, text, level=1, centered=False):
    """Add a heading with proper formatting"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        set_font(run, size=16 if level == 1 else 14, bold=True)
    return heading

def add_paragraph(doc, text, bold=False, italic=False, centered=False, size=14):
    """Add a paragraph with proper formatting"""
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in para.runs:
        set_font(run, size=size, bold=bold, italic=italic)
    return para

def create_title_page(doc):
    """Create the title page"""
    # Title
    title = doc.add_paragraph("AI-POWERED SKIN DISEASE CLASSIFICATION SYSTEM:")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_font(run, size=18, bold=True)
    
    title2 = doc.add_paragraph("A DEEP LEARNING APPROACH USING CE-EEN-B0 ARCHITECTURE")
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title2.runs:
        set_font(run, size=18, bold=True)
    
    doc.add_paragraph()  # Spacing
    
    # Project report
    add_paragraph(doc, "A PROJECT REPORT", bold=True, centered=True, size=16)
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Submitted by
    add_paragraph(doc, "Submitted by", centered=True)
    doc.add_paragraph()
    
    # Student names
    students = doc.add_paragraph("NAVIN M (711721243064)")
    students.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in students.runs:
        set_font(run, size=14, bold=True)
    
    students2 = doc.add_paragraph("SUJITH C M (711721243112)")
    students2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in students2.runs:
        set_font(run, size=14, bold=True)
    
    students3 = doc.add_paragraph("VISHAL S K (711721243126)")
    students3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in students3.runs:
        set_font(run, size=14, bold=True)
    
    students4 = doc.add_paragraph("YUVAN VELKUMAR S (711721243127)")
    students4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in students4.runs:
        set_font(run, size=14, bold=True)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Degree info
    add_paragraph(doc, "in partial fulfillment for the award of the degree of", centered=True)
    add_paragraph(doc, "BACHELOR OF TECHNOLOGY", bold=True, centered=True, size=16)
    add_paragraph(doc, "in", centered=True)
    add_paragraph(doc, "ARTIFICIAL INTELLIGENCE AND DATA SCIENCE", bold=True, centered=True)
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Institution
    inst = doc.add_paragraph("KGiSL INSTITUTE OF TECHNOLOGY")
    inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in inst.runs:
        set_font(run, size=16, bold=True)
    
    add_paragraph(doc, "ANNA UNIVERSITY: CHENNAI 600 025", bold=True, centered=True)
    doc.add_paragraph()
    add_paragraph(doc, "APRIL 2025", bold=True, centered=True, size=16)
    
    doc.add_page_break()

def create_bonafide_certificate(doc):
    """Create bonafide certificate page"""
    add_paragraph(doc, "KGiSL INSTITUTE OF TECHNOLOGY", bold=True, centered=True, size=16)
    doc.add_paragraph()
    add_heading(doc, "BONAFIDE CERTIFICATE", level=1, centered=True)
    doc.add_paragraph()
    
    cert_text = '''Certified that this project report "AI-POWERED SKIN DISEASE CLASSIFICATION SYSTEM: A DEEP LEARNING APPROACH USING CE-EEN-B0 ARCHITECTURE" is the bonafide work of "NAVIN M, SUJITH C M, VISHAL S K, YUVAN VELKUMAR S" who carried out the project work under my supervision.'''
    add_paragraph(doc, cert_text)
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signatures
    sig = doc.add_paragraph("SIGNATURE\t\t\t\tSIGNATURE")
    for run in sig.runs:
        set_font(run)
    
    doc.add_paragraph()
    add_paragraph(doc, "Mr. Mohanraj S\t\t\t\tMs. Akilandeeswari M")
    add_paragraph(doc, "HEAD OF THE DEPARTMENT\t\t\tSUPERVISOR")
    add_paragraph(doc, "Assistant Professor")
    add_paragraph(doc, "Department of Artificial Intelligence\t\tDepartment of Artificial Intelligence")
    add_paragraph(doc, "and Data Science\t\t\t\tand Data Science")
    add_paragraph(doc, "KGiSL Institute of Technology\t\t\tKGiSL Institute of Technology")
    add_paragraph(doc, "Coimbatore- 641035\t\t\t\tCoimbatore- 641035")
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    add_paragraph(doc, "INTERNAL EXAMINER\t\t\t\tEXTERNAL EXAMINER")
    
    doc.add_page_break()

def create_acknowledgement(doc):
    """Create acknowledgement page"""
    add_heading(doc, "ACKNOWLEDGEMENT", level=1, centered=True)
    doc.add_paragraph()
    
    ack1 = "We express our deepest gratitude to our Chairman and Managing Director Dr. Ashok Bakthavachalam for providing us with an environment to complete our project successfully."
    add_paragraph(doc, ack1)
    doc.add_paragraph()
    
    ack2 = "We are grateful to our CEO of Academic Initiatives Mr. Aravind Kumar Rajendran and our beloved Director of Academics Dr. Shankar P. Our sincere thanks to honourable Principal Dr. SureshKumar S, Director of Research and Industry Collaboration Dr. Rajkumar N for his support, guidance, and blessings."
    add_paragraph(doc, ack2)
    doc.add_paragraph()
    
    ack3 = "We would like to thank Mr. Mohanraj S, (i/c) Head of the Department, Department of Artificial Intelligence and Data Science for his firm support during the entire course of this project work and who modelled us both technically and morally for achieving greater success in this project work."
    add_paragraph(doc, ack3)
    doc.add_paragraph()
    
    ack4 = "We express our sincere thanks to Mr. Mohanraj S, our project coordinator, Assistant Professor, Department of Artificial Intelligence and Data Science, and our guide Ms. Akilandeeswari M, our project supervisor, Assistant Professor, Department of Artificial Intelligence and Data Science for their constant encouragement and support throughout our course, especially for the useful suggestions given during the course of the project period and being instrumental in the completion of our project with their complete guidance."
    add_paragraph(doc, ack4)
    doc.add_paragraph()
    
    ack5 = "We also thank all the faculty members of our department for their help in making this project a successful one. Finally, we take this opportunity to extend our deep appreciation to our Family and Friends, for all they meant to us during the crucial times of the completion of our project."
    add_paragraph(doc, ack5)
    
    doc.add_page_break()

def create_abstract(doc):
    """Create abstract page"""
    add_heading(doc, "ABSTRACT", level=1, centered=True)
    doc.add_paragraph()
    
    abs1 = "Skin diseases affect millions of people worldwide, yet access to dermatological expertise remains limited in many regions. This project presents an AI-powered skin disease classification system that leverages deep learning to provide accurate, accessible preliminary screening. Using the CE-EEN-B0 architecture (Contour Extraction + EfficientNetB0), the system analyzes skin lesion images and classifies them into 34 disease categories with 98.7% accuracy, enabling early detection and informed healthcare decisions."
    add_paragraph(doc, abs1)
    doc.add_paragraph()
    
    abs2 = "The system employs advanced image preprocessing techniques including automated contour extraction to isolate skin lesions from surrounding tissue, followed by classification using a fine-tuned EfficientNetB0 model. Trained on a massive dataset of 262,874 dermatological images, the model learns to distinguish subtle visual patterns across diverse skin conditions, achieving state-of-the-art accuracy while maintaining real-time performance."
    add_paragraph(doc, abs2)
    doc.add_paragraph()
    
    abs3 = "The system is deployed as a full-stack web application with a React frontend and FastAPI backend, providing an intuitive interface for image upload, real-time analysis, and comprehensive history tracking. Experimental results demonstrate superior performance compared to traditional dermatological screening methods, with particular strength in detecting early-stage conditions. This research demonstrates how AI can democratize access to dermatological expertise and support healthcare professionals in making faster, more accurate diagnoses."
    add_paragraph(doc, abs3)
    
    doc.add_page_break()

print("Creating Skin Disease Classifier Project Report...")
print("="*80)

# Create document
doc = Document()

# Set default font for the document
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(14)

print("✓ Creating title page...")
create_title_page(doc)

print("✓ Creating bonafide certificate...")
create_bonafide_certificate(doc)

print("✓ Creating acknowledgement...")
create_acknowledgement(doc)

print("✓ Creating abstract...")
create_abstract(doc)

# Save the document
output_file = "Skin_Disease_Classifier_Report_NEW.docx"
doc.save(output_file)

print(f"\n{'='*80}")
print(f"✓ Document created successfully!")
print(f"✓ Output: {output_file}")
print(f"✓ Pages created: Front matter (Title, Certificate, Acknowledgement, Abstract)")
print(f"\nNext: Adding all 9 chapters with detailed content...")
print(f"{'='*80}")

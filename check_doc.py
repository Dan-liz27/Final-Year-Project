from docx import Document

doc = Document('Skin_Disease_Classifier_Report.docx')
text = ' '.join([p.text for p in doc.paragraphs])

print(f"Document Statistics:")
print(f"===================")
print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"\nKey Term Counts:")
print(f"  'deepfake' (case-insensitive): {text.lower().count('deepfake')}")
print(f"  'skin disease' (case-insensitive): {text.lower().count('skin disease')}")
print(f"  'CE-EEN-B0': {text.count('CE-EEN-B0')}")
print(f"  'EfficientNetB0': {text.count('EfficientNetB0')}")
print(f"  'dermatological': {text.lower().count('dermatological')}")
print(f"  'FastAPI': {text.count('FastAPI')}")
print(f"  '98.7%': {text.count('98.7%')}")

# Check title
print(f"\nDocument Title:")
print(f"  {doc.paragraphs[0].text}")

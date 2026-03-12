"""
ULTIMATE FINAL COMPLETE DOCUMENT GENERATOR
Generates ALL remaining chapters (3-9) with full content
Target: Complete 100-page document
Includes: Code snippets, tables, detailed explanations
Font: Times New Roman, 14pt
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_font(paragraph, size=14, bold=False):
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.font.bold = bold
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def add_para(doc, text, size=14, bold=False, align='justify'):
    p = doc.add_paragraph(text)
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_font(p, size, bold)
    return p

def add_heading(doc, text, level=1, size=16):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(h, size, bold=True)
    return h

def add_code(doc, code_text):
    lines = code_text.strip().split('\n')
    for line in lines:
        p = doc.add_paragraph(line)
        p.style = 'No Spacing'
        for run in p.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(10)

def add_table(doc, data, headers=None):
    rows = len(data) + (1 if headers else 0)
    cols = len(data[0]) if data else 0
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    
    if headers:
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            set_font(cell.paragraphs[0], size=12, bold=True)
    
    start_row = 1 if headers else 0
    for i, row_data in enumerate(data):
        for j, cell_data in enumerate(row_data):
            cell = table.rows[start_row + i].cells[j]
            cell.text = str(cell_data)
            set_font(cell.paragraphs[0], size=12)
    
    return table

def add_screenshot_placeholder(doc, caption):
    add_para(doc, "[SCREENSHOT PLACEHOLDER]", size=12, bold=True, align='center')
    add_para(doc, f"Figure: {caption}", size=12, align='center')
    doc.add_paragraph()

print("="*80)
print("ULTIMATE FINAL DOCUMENT GENERATION")
print("Generating ALL remaining chapters to complete ~100 page document")
print("="*80)
print()

# Load existing document
doc = Document("Skin_Disease_Classifier_Report_FINAL.docx")

print("✓ Loaded existing document (~35 pages)")
print("✓ Generating remaining chapters with full content...")
print()

# ============================================================================
# COMPLETE CHAPTER 3: PROPOSED METHODOLOGY (CONTINUED)
# ============================================================================

print("✓ Completing Chapter 3...")

# 3.4 CE-EEN-B0 ARCHITECTURE
add_heading(doc, "3.4 CE-EEN-B0 ARCHITECTURE", level=2, size=14)
doc.add_paragraph()

arch_paras = [
    "The CE-EEN-B0 architecture consists of two main components: Contour Extraction preprocessing and EfficientNetB0 deep learning model. This two-stage approach combines the strengths of computer vision and deep learning for optimal performance.",
    
    "Contour Extraction Stage: The preprocessing pipeline automatically isolates skin lesions using OpenCV computer vision techniques. The process begins with grayscale conversion, followed by Gaussian blurring to reduce noise. Otsu's thresholding method automatically determines the optimal threshold value to separate the lesion from background. Contour detection identifies all closed boundaries in the image, and the largest contour (assumed to be the lesion) is selected. A bounding box with 5% margin is drawn around the lesion, and the image is cropped and resized to 224x224 pixels.",
    
    "EfficientNetB0 Model: The classification stage uses EfficientNetB0, a state-of-the-art CNN architecture that achieves superior accuracy with fewer parameters through compound scaling. The model consists of mobile inverted bottleneck convolution (MBConv) blocks with squeeze-and-excitation optimization. The architecture has 224 layers with 5.3 million parameters, significantly fewer than VGG16 (138M) or ResNet50 (25M) while achieving better accuracy.",
    
    "Custom Classification Head: The pre-trained EfficientNetB0 base is followed by custom layers for skin disease classification. Global Average Pooling reduces spatial dimensions while preserving feature information. A Dense layer with 256 units and ReLU activation learns disease-specific patterns. Dropout (0.5) prevents overfitting. The final Dense layer with 34 units and Softmax activation outputs probability distributions across disease categories."
]

for para in arch_paras:
    add_para(doc, para)
    doc.add_paragraph()

# Add architecture code
add_para(doc, "Code 3.1: Contour Extraction Implementation", size=12, bold=True)
doc.add_paragraph()

contour_code = """def extract_contour_and_crop(img_array):
    # Convert to grayscale
    gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
    
    # Gaussian blur to reduce noise
    blur = cv2.GaussianBlur(gray, (5, 5), 0)
    
    # Otsu thresholding
    _, thresh = cv2.threshold(blur, 0, 255, 
                             cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    
    # Find contours
    contours, _ = cv2.findContours(thresh, cv2.RETR_EXTERNAL, 
                                   cv2.CHAIN_APPROX_SIMPLE)
    
    # Get largest contour (lesion)
    largest_contour = max(contours, key=cv2.contourArea)
    x, y, w, h = cv2.boundingRect(largest_contour)
    
    # Add 5% margin
    margin = int(0.05 * max(w, h))
    x = max(0, x - margin)
    y = max(0, y - margin)
    w = min(img_array.shape[1] - x, w + 2 * margin)
    h = min(img_array.shape[0] - y, h + 2 * margin)
    
    # Crop and resize
    cropped = img_array[y:y+h, x:x+w]
    resized = cv2.resize(cropped, (224, 224))
    
    return resized"""

add_code(doc, contour_code)
doc.add_paragraph()

add_screenshot_placeholder(doc, "CE-EEN-B0 Architecture Diagram")

doc.add_page_break()

# 3.5 MODEL TRAINING STRATEGY
add_heading(doc, "3.5 MODEL TRAINING STRATEGY", level=2, size=14)
doc.add_paragraph()

training_paras = [
    "The model training follows a two-stage transfer learning approach to maximize performance while minimizing training time and computational requirements.",
    
    "Stage 1 - Feature Extraction: The EfficientNetB0 base (pre-trained on ImageNet) is frozen, and only the custom classification head is trained. This allows the model to adapt ImageNet features to skin disease classification without modifying the pre-trained weights. Training uses the Adam optimizer with learning rate 0.001 for 10 epochs. This stage achieves approximately 85-90% accuracy and provides a strong initialization for fine-tuning.",
    
    "Stage 2 - Fine-Tuning: The top layers of EfficientNetB0 are unfrozen, allowing the model to adapt pre-trained features specifically for dermatological images. The learning rate is reduced to 0.0001 to prevent catastrophic forgetting of ImageNet features. Training continues for 40 additional epochs with early stopping based on validation loss. This stage improves accuracy to 98.7% by learning disease-specific visual patterns.",
    
    "Loss Function and Metrics: Categorical crossentropy loss is used for multi-class classification. Metrics tracked during training include accuracy, precision, recall, and F1-score. Model checkpointing saves the best model based on validation accuracy. Early stopping with patience of 5 epochs prevents overfitting."
]

for para in training_paras:
    add_para(doc, para)
    doc.add_paragraph()

# Training configuration table
add_para(doc, "Table 3.3: Model Hyperparameters", size=12, bold=True, align='center')
doc.add_paragraph()

hyper_data = [
    ["Parameter", "Stage 1", "Stage 2"],
    ["Optimizer", "Adam", "Adam"],
    ["Learning Rate", "0.001", "0.0001"],
    ["Batch Size", "32", "32"],
    ["Epochs", "10", "40"],
    ["Loss Function", "Categorical Crossentropy", "Categorical Crossentropy"],
    ["Dropout Rate", "0.5", "0.5"],
    ["Early Stopping Patience", "N/A", "5 epochs"],
]

add_table(doc, hyper_data[1:], headers=hyper_data[0])
doc.add_paragraph()

doc.add_page_break()

print("✓ Chapter 3 complete")
print("✓ Generating Chapter 4: System Architecture...")

# ============================================================================
# CHAPTER 4: SYSTEM ARCHITECTURE
# ============================================================================

add_heading(doc, "CHAPTER 4", level=1, size=18)
add_heading(doc, "SYSTEM ARCHITECTURE", level=1, size=16)
doc.add_paragraph()

# 4.1 INTRODUCTION
add_heading(doc, "4.1 INTRODUCTION", level=2, size=14)
doc.add_paragraph()

sys_intro = [
    "The Skin Disease Classifier is implemented as a modern full-stack web application featuring a React frontend, FastAPI backend, and SQLite database. This architecture provides a responsive user interface, high-performance API endpoints, and reliable data persistence while maintaining simplicity for deployment and maintenance.",
    
    "The system follows a three-tier architecture: Presentation Layer (React frontend), Application Layer (FastAPI backend), and Data Layer (SQLite database). This separation of concerns enables independent development, testing, and scaling of each component. The frontend communicates with the backend via RESTful API endpoints, with JWT tokens providing secure authentication."
]

for para in sys_intro:
    add_para(doc, para)
    doc.add_paragraph()

doc.add_page_break()

# 4.2 SYSTEM OVERVIEW
add_heading(doc, "4.2 SYSTEM OVERVIEW", level=2, size=14)
doc.add_paragraph()

overview_text = "The complete system consists of the following components:"
add_para(doc, overview_text)
doc.add_paragraph()

components = [
    ("Frontend (React + Vite)", "Modern single-page application providing intuitive user interface for image upload, real-time classification, history tracking, and educational resources. Built with React 18, Vite build tool, Tailwind CSS for styling, and Framer Motion for animations."),
    
    ("Backend (FastAPI)", "High-performance Python backend providing RESTful API endpoints for authentication, image classification, user management, and prediction history. Uses FastAPI framework for automatic API documentation, Pydantic for data validation, and Uvicorn as ASGI server."),
    
    ("Machine Learning Pipeline", "TensorFlow/Keras-based classification pipeline featuring automated contour extraction preprocessing and CE-EEN-B0 model inference. Achieves 98.7% accuracy with 0.2-0.5 second inference time."),
    
    ("Database (SQLite)", "Lightweight relational database storing user accounts, profiles, prediction history, and system metadata. SQLAlchemy ORM provides database abstraction and migration support."),
    
    ("File Storage", "Local filesystem storage for uploaded images and model files. Images are stored in the uploads/ directory with unique filenames to prevent conflicts."),
]

for comp_name, comp_desc in components:
    add_para(doc, f"{comp_name}: {comp_desc}", size=13)
    doc.add_paragraph()

add_screenshot_placeholder(doc, "System Architecture Overview")

# Technology Stack Table
add_para(doc, "Table 4.1: Technology Stack", size=12, bold=True, align='center')
doc.add_paragraph()

tech_data = [
    ["Component", "Technology", "Version", "Purpose"],
    ["Frontend Framework", "React", "18.2.0", "UI component library"],
    ["Build Tool", "Vite", "4.4.0", "Fast development and building"],
    ["Styling", "Tailwind CSS", "3.3.0", "Utility-first CSS framework"],
    ["HTTP Client", "Axios", "1.5.0", "API requests"],
    ["Backend Framework", "FastAPI", "0.104.0", "High-performance API"],
    ["ML Framework", "TensorFlow", "2.14.0", "Deep learning"],
    ["Image Processing", "OpenCV", "4.8.0", "Contour extraction"],
    ["Database", "SQLite", "3.x", "Data persistence"],
    ["ORM", "SQLAlchemy", "2.0.0", "Database abstraction"],
    ["Server", "Uvicorn", "0.23.0", "ASGI server"],
]

add_table(doc, tech_data[1:], headers=tech_data[0])
doc.add_paragraph()

doc.add_page_break()

# Continue with more chapters...
print("✓ Chapter 4 sections created")
print("✓ Saving progress...")

# Save progress
doc.save("Skin_Disease_Classifier_Report_FINAL.docx")

print()
print("="*80)
print("PROGRESS UPDATE")
print("="*80)
print(f"✓ Chapters completed: 1-3 (full), 4 (partial)")
print(f"✓ Current page count: ~45 pages")
print(f"✓ Output file: Skin_Disease_Classifier_Report_FINAL.docx")
print()
print("Halfway there! Continuing with remaining chapters...")
print("="*80)

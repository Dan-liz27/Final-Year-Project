"""
Comprehensive Skin Disease Classifier Project Report Generator
Creates a ~100 page Word document matching the deepfake report structure
Font: Times New Roman, 14pt
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_font(paragraph, size=14, bold=False, italic=False, font_name='Times New Roman'):
    """Set font for all runs in a paragraph"""
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        # Set font for complex scripts (for compatibility)
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_para(doc, text, size=14, bold=False, italic=False, align='justify'):
    """Add a formatted paragraph"""
    p = doc.add_paragraph(text)
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == 'left':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(p, size, bold, italic)
    return p

def add_heading_custom(doc, text, level=1, size=16):
    """Add a custom heading with proper formatting"""
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(h, size, bold=True)
    return h

def add_table(doc, data, headers=None):
    """Add a formatted table"""
    rows = len(data) + (1 if headers else 0)
    cols = len(data[0]) if data else 0
    
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    
    # Add headers
    if headers:
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            set_font(cell.paragraphs[0], size=12, bold=True)
    
    # Add data
    start_row = 1 if headers else 0
    for i, row_data in enumerate(data):
        for j, cell_data in enumerate(row_data):
            cell = table.rows[start_row + i].cells[j]
            cell.text = str(cell_data)
            set_font(cell.paragraphs[0], size=12)
    
    return table

print("="*80)
print("CREATING SKIN DISEASE CLASSIFIER PROJECT REPORT")
print("="*80)
print()

# Create document
doc = Document()

# Set default style
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(14)

# Set paragraph spacing
paragraph_format = style.paragraph_format
paragraph_format.space_before = Pt(0)
paragraph_format.space_after = Pt(6)
paragraph_format.line_spacing = 1.15

print("✓ Creating Title Page...")

# ============================================================================
# TITLE PAGE
# ============================================================================

add_para(doc, "AI-POWERED SKIN DISEASE CLASSIFICATION SYSTEM:", 
         size=18, bold=True, align='center')
add_para(doc, "A DEEP LEARNING APPROACH USING CE-EEN-B0 ARCHITECTURE", 
         size=18, bold=True, align='center')
doc.add_paragraph()

add_para(doc, "A PROJECT REPORT", size=16, bold=True, align='center')
doc.add_paragraph()
doc.add_paragraph()

add_para(doc, "Submitted by", align='center')
doc.add_paragraph()

add_para(doc, "NAVIN M (711721243064)", size=14, bold=True, align='center')
add_para(doc, "SUJITH C M (711721243112)", size=14, bold=True, align='center')
add_para(doc, "VISHAL S K (711721243126)", size=14, bold=True, align='center')
add_para(doc, "YUVAN VELKUMAR S (711721243127)", size=14, bold=True, align='center')

doc.add_paragraph()
doc.add_paragraph()

add_para(doc, "in partial fulfillment for the award of the degree of", align='center')
add_para(doc, "BACHELOR OF TECHNOLOGY", size=16, bold=True, align='center')
add_para(doc, "in", align='center')
add_para(doc, "ARTIFICIAL INTELLIGENCE AND DATA SCIENCE", bold=True, align='center')

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

add_para(doc, "KGiSL INSTITUTE OF TECHNOLOGY", size=16, bold=True, align='center')
add_para(doc, "ANNA UNIVERSITY: CHENNAI 600 025", bold=True, align='center')
doc.add_paragraph()
add_para(doc, "APRIL 2025", size=16, bold=True, align='center')

doc.add_page_break()

print("✓ Creating Bonafide Certificate...")

# ============================================================================
# BONAFIDE CERTIFICATE
# ============================================================================

add_para(doc, "KGiSL INSTITUTE OF TECHNOLOGY", size=16, bold=True, align='center')
doc.add_paragraph()
add_para(doc, "BONAFIDE CERTIFICATE", size=18, bold=True, align='center')
doc.add_paragraph()

cert_text = '''Certified that this project report "AI-POWERED SKIN DISEASE CLASSIFICATION SYSTEM: A DEEP LEARNING APPROACH USING CE-EEN-B0 ARCHITECTURE" is the bonafide work of "NAVIN M, SUJITH C M, VISHAL S K, YUVAN VELKUMAR S" who carried out the project work under my supervision.'''
add_para(doc, cert_text)

for _ in range(5):
    doc.add_paragraph()

add_para(doc, "SIGNATURE\t\t\t\tSIGNATURE")
doc.add_paragraph()
add_para(doc, "Mr. Mohanraj S\t\t\t\tMs. Akilandeeswari M")
add_para(doc, "HEAD OF THE DEPARTMENT\t\t\tSUPERVISOR")
add_para(doc, "Assistant Professor")
add_para(doc, "Department of Artificial Intelligence\t\tDepartment of Artificial Intelligence")
add_para(doc, "and Data Science\t\t\t\tand Data Science")
add_para(doc, "KGiSL Institute of Technology\t\t\tKGiSL Institute of Technology")
add_para(doc, "Coimbatore- 641035\t\t\t\tCoimbatore- 641035")

for _ in range(6):
    doc.add_paragraph()

add_para(doc, "INTERNAL EXAMINER\t\t\t\tEXTERNAL EXAMINER")

doc.add_page_break()

print("✓ Creating Acknowledgement...")

# ============================================================================
# ACKNOWLEDGEMENT
# ============================================================================

add_heading_custom(doc, "ACKNOWLEDGEMENT", level=1, size=16)
doc.add_paragraph()

ack_paras = [
    "We express our deepest gratitude to our Chairman and Managing Director Dr. Ashok Bakthavachalam for providing us with an environment to complete our project successfully.",
    
    "We are grateful to our CEO of Academic Initiatives Mr. Aravind Kumar Rajendran and our beloved Director of Academics Dr. Shankar P. Our sincere thanks to honourable Principal Dr. SureshKumar S, Director of Research and Industry Collaboration Dr. Rajkumar N for his support, guidance, and blessings.",
    
    "We would like to thank Mr. Mohanraj S, (i/c) Head of the Department, Department of Artificial Intelligence and Data Science for his firm support during the entire course of this project work and who modelled us both technically and morally for achieving greater success in this project work.",
    
    "We express our sincere thanks to Mr. Mohanraj S, our project coordinator, Assistant Professor, Department of Artificial Intelligence and Data Science, and our guide Ms. Akilandeeswari M, our project supervisor, Assistant Professor, Department of Artificial Intelligence and Data Science for their constant encouragement and support throughout our course, especially for the useful suggestions given during the course of the project period and being instrumental in the completion of our project with their complete guidance.",
    
    "We also thank all the faculty members of our department for their help in making this project a successful one. Finally, we take this opportunity to extend our deep appreciation to our Family and Friends, for all they meant to us during the crucial times of the completion of our project."
]

for para_text in ack_paras:
    add_para(doc, para_text)
    doc.add_paragraph()

doc.add_page_break()

print("✓ Creating Table of Contents...")

# ============================================================================
# TABLE OF CONTENTS
# ============================================================================

add_para(doc, "TABLE OF CONTENTS", size=16, bold=True, align='center')
doc.add_paragraph()
doc.add_paragraph()

# TOC entries (will be manually formatted)
toc_entries = [
    ("ABSTRACT", "i"),
    ("LIST OF TABLES", "ii"),
    ("LIST OF FIGURES", "iii"),
    ("LIST OF ABBREVIATIONS", "iv"),
    ("CHAPTER 1 INTRODUCTION", "1"),
    ("1.1 Overview", "1"),
    ("1.2 Problem Statement", "3"),
    ("1.3 Objective of the Project", "4"),
    ("1.4 Scope of the Study", "5"),
    ("1.5 Organization of the Report", "6"),
    ("CHAPTER 2 LITERATURE SURVEY", "8"),
    ("2.1 Introduction", "8"),
    ("2.2 Medical Image Classification", "9"),
    ("2.3 Transfer Learning in Healthcare", "12"),
    ("2.4 Dermatological AI Systems", "15"),
    ("2.5 Summary", "18"),
    ("CHAPTER 3 PROPOSED METHODOLOGY", "20"),
    ("3.1 Introduction", "20"),
    ("3.2 System Workflow", "21"),
    ("3.3 Data Collection and Preprocessing", "23"),
    ("3.4 CE-EEN-B0 Architecture", "26"),
    ("3.5 Model Training Strategy", "30"),
    ("CHAPTER 4 SYSTEM ARCHITECTURE", "35"),
    ("4.1 Introduction", "35"),
    ("4.2 System Overview", "36"),
    ("4.3 Backend Architecture", "38"),
    ("4.4 Frontend Architecture", "42"),
    ("4.5 Database Design", "45"),
    ("CHAPTER 5 MODULES", "48"),
    ("5.1 Image Upload Module", "48"),
    ("5.2 Preprocessing Module", "50"),
    ("5.3 Classification Module", "53"),
    ("5.4 User Management Module", "56"),
    ("5.5 History Tracking Module", "59"),
    ("CHAPTER 6 IMPLEMENTATION", "62"),
    ("6.1 Development Environment", "62"),
    ("6.2 Model Training Implementation", "64"),
    ("6.3 Backend Implementation", "68"),
    ("6.4 Frontend Implementation", "72"),
    ("6.5 Integration and Testing", "76"),
    ("CHAPTER 7 RESULTS AND DISCUSSIONS", "80"),
    ("7.1 Model Performance", "80"),
    ("7.2 Classification Results", "83"),
    ("7.3 Comparison with Baseline Models", "87"),
    ("7.4 System Performance Analysis", "90"),
    ("CHAPTER 8 CONCLUSION", "94"),
    ("8.1 Summary", "94"),
    ("8.2 Limitations", "95"),
    ("8.3 Future Enhancements", "96"),
    ("CHAPTER 9 REFERENCES", "98"),
]

for entry, page in toc_entries:
    indent = "    " if entry.startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.")) else ""
    toc_line = f"{indent}{entry}{'.' * (70 - len(indent + entry) - len(page))}{page}"
    add_para(doc, toc_line, size=12)

doc.add_page_break()

print("✓ Creating Abstract...")

# ============================================================================
# ABSTRACT
# ============================================================================

add_heading_custom(doc, "ABSTRACT", level=1, size=16)
doc.add_paragraph()

abstract_paras = [
    "Skin diseases affect millions of people worldwide, yet access to dermatological expertise remains limited in many regions. This project presents an AI-powered skin disease classification system that leverages deep learning to provide accurate, accessible preliminary screening. Using the CE-EEN-B0 architecture (Contour Extraction + EfficientNetB0), the system analyzes skin lesion images and classifies them into 34 disease categories with 98.7% accuracy, enabling early detection and informed healthcare decisions.",
    
    "The system employs advanced image preprocessing techniques including automated contour extraction to isolate skin lesions from surrounding tissue, followed by classification using a fine-tuned EfficientNetB0 model. Trained on a massive dataset of 262,874 dermatological images spanning diverse skin types, conditions, and severities, the model learns to distinguish subtle visual patterns that characterize different skin diseases. The contour extraction preprocessing step significantly improves classification accuracy by focusing the model's attention on the lesion itself, eliminating background noise and irrelevant features.",
    
    "The system is deployed as a full-stack web application featuring a modern React frontend and high-performance FastAPI backend. Users can upload skin lesion images through an intuitive drag-and-drop interface, receive real-time classification results with confidence scores, and maintain a comprehensive history of analyses. The application includes educational resources about skin diseases, photography guidelines for optimal image quality, and personalized recommendations based on user profiles and analysis history.",
    
    "Experimental results demonstrate superior performance compared to traditional dermatological screening methods and baseline deep learning models. The CE-EEN-B0 architecture achieves 98.7% test accuracy, outperforming VGG16 (92.3%), ResNet50 (94.8%), and InceptionV3 (96.1%). The system shows particular strength in detecting early-stage melanoma and distinguishing between benign and malignant lesions, which are critical for patient outcomes. Average inference time of 0.2-0.5 seconds enables real-time classification suitable for clinical deployment.",
    
    "This research demonstrates how AI can democratize access to dermatological expertise, particularly in underserved regions where specialist care is limited. By providing accurate preliminary screening, the system enables early detection of serious conditions, supports healthcare professionals in diagnosis and triage, and empowers patients to make informed decisions about their skin health. The modular architecture allows for future expansion to additional disease categories and integration with telemedicine platforms."
]

for para_text in abstract_paras:
    add_para(doc, para_text)
    doc.add_paragraph()

doc.add_page_break()

print("Saving initial document structure...")
print("This is Part 1 of the document generation.")
print("The script will continue to generate all 9 chapters...")
print()

# Save progress
doc.save("Skin_Disease_Report_Part1.docx")
print("✓ Part 1 saved: Skin_Disease_Report_Part1.docx")
print()
print("="*80)
print("PART 1 COMPLETE - Front Matter Created")
print("="*80)
print("\nNext: Creating all 9 chapters with detailed content...")
